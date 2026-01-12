import re
import io
import json
import time
import uuid
import hashlib
import secrets
from dataclasses import dataclass
from typing import List, Tuple, Dict, Any, Optional
from datetime import datetime, timedelta, timezone

import streamlit as st
from pypdf import PdfReader
from docx import Document as DocxDocument

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas as pdf_canvas

from supabase import create_client

# =========================
# App Config
# =========================
APP_NAME = "DocuLite"
APP_VERSION = "MASTER v4.0 (Docs + Slides PDF + Sections + Glossary + Flashcards + Compare + Export + Supabase PINs)"
st.set_page_config(page_title=APP_NAME, page_icon="📄", layout="wide")


# =========================
# Tier Limits (your rules)
# =========================
@dataclass
class Tier:
    name: str
    max_docs: int
    max_pages: int
    max_mb: int
    max_devices_default: int

FREE_TIER = Tier("FREE", max_docs=2, max_pages=150, max_mb=300, max_devices_default=1)
PAID_TIER = Tier("PAID", max_docs=10, max_pages=200, max_mb=300, max_devices_default=2)
OWNER_TIER = Tier("OWNER", max_docs=9999, max_pages=9999, max_mb=300, max_devices_default=5)


# =========================
# Secrets / Supabase
# =========================
def _secret(name: str, default: str = "") -> str:
    try:
        return str(st.secrets.get(name, default)).strip()
    except Exception:
        return default

SUPABASE_URL = _secret("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = _secret("SUPABASE_SERVICE_ROLE_KEY")
PIN_SALT = _secret("PIN_SALT")
OWNER_ADMIN_PIN = _secret("OWNER_ADMIN_PIN")

@st.cache_resource
def sb():
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        raise RuntimeError("Supabase secrets missing. Set SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY in Streamlit Secrets.")
    return create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)

def hash_pin(pin: str) -> str:
    s = (PIN_SALT + (pin or "").strip()).encode("utf-8")
    return hashlib.sha256(s).hexdigest()

def generate_pin() -> str:
    raw = secrets.token_urlsafe(10).replace("-", "").replace("_", "").upper()
    raw = raw[:12]
    return f"{raw[:4]}-{raw[4:8]}-{raw[8:12]}"

def get_headers_safe() -> Dict[str, str]:
    # Streamlit may or may not expose headers depending on environment
    try:
        return dict(st.context.headers)  # type: ignore
    except Exception:
        return {}

def device_fingerprint() -> str:
    """
    Not perfect, but blocks casual PIN sharing.
    We also store a session UUID to differentiate devices behind same NAT.
    """
    if "device_uuid" not in st.session_state:
        st.session_state.device_uuid = uuid.uuid4().hex

    h = get_headers_safe()
    ua = h.get("User-Agent", "")
    ip = h.get("X-Forwarded-For", "")
    ip = ip.split(",")[0].strip()

    base = f"{ua}|{ip}|{st.session_state.device_uuid}"
    return hashlib.sha256((PIN_SALT + base).encode("utf-8")).hexdigest()[:24]

def tier_from_name(name: str) -> Tier:
    if name == "OWNER":
        return OWNER_TIER
    if name == "PAID":
        return PAID_TIER
    return FREE_TIER

def verify_pin_and_bind(pin: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    pin = (pin or "").strip()
    if not pin:
        return False, "Enter a PIN.", None

    if not PIN_SALT:
        return False, "PIN_SALT is missing in Secrets.", None

    h = hash_pin(pin)
    client = sb()

    res = client.table("pins").select("*").eq("pin_hash", h).limit(1).execute()
    if not res.data:
        return False, "Invalid PIN.", None

    row = res.data[0]
    if row.get("status") != "ACTIVE":
        return False, "This PIN is revoked.", row

    if row.get("expires_at"):
        exp = datetime.fromisoformat(str(row["expires_at"]).replace("Z", "+00:00"))
        if datetime.now(timezone.utc) > exp:
            return False, "This PIN has expired.", row

    dev = device_fingerprint()
    devices = row.get("devices") or []
    max_devices = int(row.get("max_devices") or 1)

    if dev not in devices:
        if len(devices) >= max_devices:
            return False, "PIN device limit reached. Contact the owner.", row
        devices.append(dev)
        client.table("pins").update({"devices": devices}).eq("id", row["id"]).execute()
        row["devices"] = devices

    return True, f"Unlocked: {row['tier']}", row


# =========================
# Text Cleanup (fix spaced letters)
# =========================
STOPWORDS = set("""
a an the and or but if while to of in on for with without at by from as is are was were be been being
this that these those it its into about over under between among than then so because therefore
i you he she they we them us our your my mine theirs his her
""".split())

def fix_spaced_letters(text: str) -> str:
    if not text:
        return text

    def _join(match):
        return match.group(0).replace(" ", "")

    # "C O N T E N T S" -> "CONTENTS"
    text = re.sub(r'(?<!\w)(?:[A-Za-z]\s){2,}[A-Za-z](?!\w)', _join, text)

    # hyphen line break: "inter-\nnational" -> "international"
    text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)

    # tidy
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

def normalize_text(s: str) -> str:
    if not s:
        return ""
    s = s.replace("\r", "\n").replace("\u00a0", " ")
    s = re.sub(r"[ \t]{2,}", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()
    s = fix_spaced_letters(s)
    # convert single newlines into spaces (keep paragraphs)
    s = re.sub(r"(?<!\n)\n(?!\n)", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()
    return s

def sentence_split(text: str) -> List[str]:
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if not text:
        return []
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z0-9"“‘])', text)
    return [p.strip() for p in parts if p.strip()]

def tokenize_words(text: str) -> List[str]:
    words = re.findall(r"[A-Za-z][A-Za-z'\-]{1,}", (text or "").lower())
    return [w for w in words if w not in STOPWORDS and len(w) >= 3]

def jaccard(a: str, b: str) -> float:
    sa = set(tokenize_words(a))
    sb = set(tokenize_words(b))
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / max(1, len(sa | sb))


# =========================
# Keyword + Summary (fast, no NLTK, no sklearn)
# =========================
def extract_keywords(text: str, top_k: int = 12) -> List[str]:
    """
    A fast “RAKE-lite” scoring:
    - unigram + bigram candidates
    - frequency * log-length
    """
    text = (text or "")
    if len(text) < 60:
        return []

    words = tokenize_words(text)
    if not words:
        return []

    freq: Dict[str, int] = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1

    # bigrams
    bigrams: Dict[str, int] = {}
    for i in range(len(words) - 1):
        bg = f"{words[i]} {words[i+1]}"
        if words[i] in STOPWORDS or words[i+1] in STOPWORDS:
            continue
        bigrams[bg] = bigrams.get(bg, 0) + 1

    scored: List[Tuple[str, float]] = []
    for w, c in freq.items():
        scored.append((w, c * 1.0))

    for bg, c in bigrams.items():
        scored.append((bg, c * 1.8))

    scored.sort(key=lambda x: x[1], reverse=True)

    out = []
    seen = set()
    for term, _ in scored:
        t = term.strip().lower()
        if t in seen:
            continue
        seen.add(t)
        out.append(term)
        if len(out) >= top_k:
            break
    return out

def summarize(text: str, n_sent: int, user_type: str) -> List[str]:
    sents = sentence_split(text)
    if not sents:
        return []

    # user_type tuning
    if user_type == "Primary learner":
        n_sent = max(3, min(n_sent, 5))
    elif user_type == "High school learner":
        n_sent = max(4, min(n_sent, 7))
    elif user_type == "Varsity/College":
        n_sent = max(6, min(n_sent, 10))
    elif user_type == "Trader":
        n_sent = max(6, min(n_sent, 10))
    else:
        n_sent = max(5, min(n_sent, 9))

    kws = set([k.split()[0].lower() for k in extract_keywords(text, top_k=20) if k])
    scores = []
    for idx, s in enumerate(sents):
        w = tokenize_words(s)
        hit = sum(1 for x in w if x in kws)
        density = hit / max(1, len(w))
        length_penalty = 1.0 if len(s) <= 220 else max(0.6, 220 / len(s))
        score = (hit + 2.0 * density) * length_penalty
        scores.append((idx, s, score))

    scores.sort(key=lambda x: x[2], reverse=True)
    picked = sorted(scores[: min(n_sent, len(scores))], key=lambda x: x[0])
    return [re.sub(r"\s+", " ", s).strip() for _, s, _ in picked]

def key_takeaways(text: str, k: int = 6) -> List[str]:
    sents = sentence_split(text)
    if not sents:
        return []
    # take top summary sentences then diversify
    candidates = summarize(text, n_sent=min(14, max(8, k * 2)), user_type="Varsity/College")
    out: List[str] = []
    for s in candidates:
        if len(out) >= k:
            break
        if all(jaccard(s, prev) < 0.55 for prev in out):
            out.append(s)
    return out


# =========================
# Questions (non-repeating + difficulty by user type)
# =========================
def generate_questions(text: str, user_type: str, n: int = 10) -> List[str]:
    kws = extract_keywords(text, top_k=max(20, n * 3))
    takes = key_takeaways(text, k=5)

    if user_type == "Primary learner":
        templates = [
            "What does '{k}' mean?",
            "Give an example of '{k}'.",
            "Explain '{k}' in your own words.",
            "Write one fact about '{k}'.",
        ]
    elif user_type == "High school learner":
        templates = [
            "Define '{k}'.",
            "Why is '{k}' important in this document?",
            "Compare '{k}' with another idea from the text.",
            "Give an example and explain '{k}'.",
        ]
    elif user_type == "Trader":
        templates = [
            "Explain '{k}' and how to use it in a trading plan (steps).",
            "What mistake happens if '{k}' is misunderstood?",
            "Create a checklist that includes '{k}'.",
            "When might '{k}' fail? Give one condition.",
        ]
    else:  # Varsity/General
        templates = [
            "Critically evaluate '{k}' based on the text.",
            "How does '{k}' connect to the main argument?",
            "Identify assumptions behind '{k}'.",
            "Apply '{k}' to a real-world scenario and justify.",
        ]

    base: List[str] = []
    for i, k in enumerate(kws):
        base.append(templates[i % len(templates)].format(k=k))

    for t in takes:
        base.append(f"Explain the meaning of this statement and why it matters: “{t[:120]}…”")

    # Diversify so questions don’t repeat
    picked: List[str] = []
    for q in base:
        if len(picked) >= n:
            break
        if all(jaccard(q, prev) < 0.55 for prev in picked):
            picked.append(q)
    return picked[:n]


# =========================
# Highlighting (in original text)
# =========================
def highlight_html(text: str, keywords: List[str], max_chars: int = 120000) -> str:
    text = (text or "")[:max_chars]
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    keywords = sorted([k.strip() for k in (keywords or []) if k.strip()], key=len, reverse=True)

    for kw in keywords[:25]:
        pat = re.compile(rf"(?i)\b({re.escape(kw)})\b")
        safe = pat.sub(r"<mark>\1</mark>", safe)

    return f"<div style='white-space:pre-wrap; line-height:1.65'>{safe}</div>"


# =========================
# PDF / DOCX / TXT Extraction
# =========================
def bytes_mb(b: bytes) -> float:
    return len(b) / (1024 * 1024)

def extract_docx_text(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return normalize_text("\n".join(paras))

def extract_txt_text(file_bytes: bytes) -> str:
    try:
        return normalize_text(file_bytes.decode("utf-8", errors="ignore"))
    except Exception:
        return ""

def pdf_total_pages(file_bytes: bytes) -> int:
    reader = PdfReader(io.BytesIO(file_bytes))
    return len(reader.pages)

def extract_pdf_range(file_bytes: bytes, start_page1: int, end_page1: int, progress_cb=None) -> Tuple[str, List[str], int]:
    """
    Returns:
      full_text, per_page_texts, pages_read
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    total = len(reader.pages)
    start0 = max(0, min(total - 1, start_page1 - 1))
    end0 = max(0, min(total - 1, end_page1 - 1))
    if end0 < start0:
        start0, end0 = end0, start0

    pages_to_read = end0 - start0 + 1
    per_page = []
    out = []

    for j, i in enumerate(range(start0, end0 + 1), start=1):
        try:
            t = reader.pages[i].extract_text() or ""
        except Exception:
            t = ""
        t = normalize_text(t)
        per_page.append(t)
        if t:
            out.append(t)
        if progress_cb:
            progress_cb(j, pages_to_read)

    return normalize_text("\n\n".join(out)), per_page, pages_to_read


# =========================
# TOC/Headings Section Detection (PDF)
# =========================
HEADING_PATTERNS = [
    r'^\s*chapter\s+\d+\b.*$',
    r'^\s*\d+(\.\d+)*\s+[A-Z].*$',
    r'^\s*[A-Z][A-Z0-9\s:,-]{6,}$',
]

def looks_like_heading(line: str) -> bool:
    s = (line or "").strip()
    if len(s) < 6 or len(s) > 120:
        return False
    for pat in HEADING_PATTERNS:
        if re.match(pat, s, flags=re.IGNORECASE):
            return True
    return False

def _flatten_outline(outline: Any) -> List[Any]:
    items = []
    if outline is None:
        return items
    if isinstance(outline, list):
        for x in outline:
            items.extend(_flatten_outline(x))
    else:
        items.append(outline)
    return items

def toc_sections(file_bytes: bytes) -> List[Tuple[str, int]]:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        outline = getattr(reader, "outline", None)
        if not outline:
            return []
        flat = _flatten_outline(outline)
        found: List[Tuple[str, int]] = []
        seen = set()
        for item in flat:
            try:
                title = getattr(item, "title", None)
                if not title:
                    continue
                page0 = reader.get_destination_page_number(item)
                if page0 is None:
                    continue
                key = re.sub(r"\s+", " ", str(title)).lower()
                if key in seen:
                    continue
                seen.add(key)
                found.append((str(title).strip()[:80], int(page0)))
            except Exception:
                continue
        found.sort(key=lambda x: x[1])
        if found and found[0][1] > 0:
            found.insert(0, ("Start", 0))
        return found
    except Exception:
        return []

def scan_headings(file_bytes: bytes, scan_pages: int = 30) -> List[Tuple[str, int]]:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        total = len(reader.pages)
        scan_pages = min(max(5, scan_pages), total)
        hits: List[Tuple[str, int]] = []
        seen = set()
        for i in range(scan_pages):
            try:
                txt = reader.pages[i].extract_text() or ""
            except Exception:
                txt = ""
            txt = normalize_text(txt)
            for ln in txt.splitlines()[:40]:
                t = ln.strip()
                if not t:
                    continue
                if looks_like_heading(t):
                    key = re.sub(r"\s+", " ", t).lower()
                    if key in seen:
                        continue
                    seen.add(key)
                    hits.append((t[:80], i))
        if hits and hits[0][1] > 0:
            hits.insert(0, ("Start", 0))
        hits.sort(key=lambda x: x[1])
        return hits[:50]
    except Exception:
        return []

def heading_ranges(headings: List[Tuple[str, int]], total_pages: int) -> List[Tuple[str, int, int]]:
    ranges = []
    for idx, (title, start0) in enumerate(headings):
        start1 = start0 + 1
        if idx + 1 < len(headings):
            next_start0 = headings[idx + 1][1]
            end1 = max(start1, next_start0)
        else:
            end1 = total_pages
        end1 = max(start1, end1)
        ranges.append((title, start1, end1))
    return ranges


# =========================
# Sections Splitter (text-based, for non-PDF or fallback)
# =========================
def split_into_sections(text: str) -> List[Tuple[str, str]]:
    lines = (text or "").splitlines()
    sections: List[Tuple[str, str]] = []
    current_title = "Start"
    buf: List[str] = []
    for ln in lines:
        if looks_like_heading(ln):
            chunk = "\n".join(buf).strip()
            if chunk:
                sections.append((current_title, chunk))
            current_title = ln.strip()
            buf = []
        else:
            buf.append(ln)
    final = "\n".join(buf).strip()
    if final:
        sections.append((current_title, final))
    if len(sections) <= 1:
        return [("Whole Document", (text or "").strip())]
    return sections


# =========================
# Export (PDF + DOCX)
# =========================
def export_pdf(title: str, summary: List[str], takeaways: List[str], questions: List[str]) -> bytes:
    buf = io.BytesIO()
    c = pdf_canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x = 2.0 * cm
    y = height - 2.0 * cm

    def draw_wrapped(txt: str, y: float, font="Helvetica", size=11, leading=14) -> float:
        c.setFont(font, size)
        max_w = width - 4.0 * cm
        words = txt.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, font, size) <= max_w:
                line = test
            else:
                c.drawString(x, y, line)
                y -= leading
                line = w
                if y < 2.0 * cm:
                    c.showPage()
                    y = height - 2.0 * cm
                    c.setFont(font, size)
        if line:
            c.drawString(x, y, line)
            y -= leading
        return y

    c.setFont("Helvetica-Bold", 16)
    c.drawString(x, y, title)
    y -= 24

    c.setFont("Helvetica", 10)
    y = draw_wrapped(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", y, size=10)
    y -= 8

    c.setFont("Helvetica-Bold", 13)
    c.drawString(x, y, "Summary")
    y -= 18
    for s in summary or ["(No summary)"]:
        y = draw_wrapped(f"- {s}", y)

    y -= 10
    c.setFont("Helvetica-Bold", 13)
    c.drawString(x, y, "Key Takeaways")
    y -= 18
    for i, t in enumerate(takeaways or ["(No takeaways)"], start=1):
        y = draw_wrapped(f"{i}. {t}", y)

    y -= 10
    c.setFont("Helvetica-Bold", 13)
    c.drawString(x, y, "Questions")
    y -= 18
    for i, q in enumerate(questions or ["(No questions)"], start=1):
        y = draw_wrapped(f"{i}. {q}", y)

    c.save()
    buf.seek(0)
    return buf.getvalue()

def export_docx(title: str, summary: List[str], takeaways: List[str], questions: List[str]) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading("Summary", level=2)
    for s in summary or ["(No summary)"]:
        doc.add_paragraph(f"- {s}")
    doc.add_heading("Key Takeaways", level=2)
    for i, t in enumerate(takeaways or ["(No takeaways)"], start=1):
        doc.add_paragraph(f"{i}. {t}")
    doc.add_heading("Questions", level=2)
    for i, q in enumerate(questions or ["(No questions)"], start=1):
        doc.add_paragraph(f"{i}. {q}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# =========================
# Flashcards
# =========================
def make_flashcards(text: str, keywords: List[str], n_cards: int = 10) -> List[Dict[str, Any]]:
    sents = [s for s in sentence_split(text) if len(s) >= 60]
    if not sents:
        sents = sentence_split(text)

    pool = [k for k in (keywords or []) if 3 <= len(k) <= 35]
    pool = list(dict.fromkeys(pool))

    cards = []
    used = set()

    for sent in sents:
        if len(cards) >= n_cards:
            break
        chosen = None
        for k in pool:
            if re.search(rf"(?i)\b{re.escape(k)}\b", sent):
                chosen = k
                break
        if not chosen:
            continue

        key = (sent[:80] + chosen).lower()
        if key in used:
            continue
        used.add(key)

        blanked = re.sub(rf"(?i)\b{re.escape(chosen)}\b", "_____", sent, count=1)

        # alternate typed and mcq
        if len(cards) % 2 == 0:
            cards.append({"type": "typed", "q": f"Fill in the missing term:\n\n{blanked}", "a": chosen})
        else:
            distractors = [x for x in pool if x.lower() != chosen.lower()]
            distractors = distractors[:3] if len(distractors) >= 3 else distractors
            options = [chosen] + distractors
            options = list(dict.fromkeys(options))
            cards.append({"type": "mcq", "q": f"Which term best completes the sentence?\n\n{blanked}", "a": chosen, "options": options})

    return cards


# =========================
# Glossary (search word in text)
# =========================
def glossary_search(term: str, text: str, limit: int = 6) -> List[str]:
    term = (term or "").strip()
    if not term:
        return []
    sents = sentence_split(text)
    hits = []
    for s in sents:
        if re.search(rf"(?i)\b{re.escape(term)}\b", s):
            hits.append(s.strip())
        if len(hits) >= limit:
            break
    return hits


# =========================
# Session State
# =========================
def init_state():
    if "auth_ok" not in st.session_state:
        st.session_state.auth_ok = False
    if "auth_row" not in st.session_state:
        st.session_state.auth_row = None
    if "tier" not in st.session_state:
        st.session_state.tier = FREE_TIER
    if "docs_used" not in st.session_state:
        st.session_state.docs_used = 0
    if "history" not in st.session_state:
        st.session_state.history = []
    if "last_result" not in st.session_state:
        st.session_state.last_result = None
    if "cards" not in st.session_state:
        st.session_state.cards = []
    if "card_answers" not in st.session_state:
        st.session_state.card_answers = {}

init_state()


# =========================
# UI Header
# =========================
st.title(f"📄 {APP_NAME}")
st.caption(APP_VERSION)


# =========================
# Sidebar: Login + Admin
# =========================
with st.sidebar:
    st.subheader("🔐 Login (Unique PINs via Supabase)")
    pin_input = st.text_input("Enter PIN", type="password", placeholder="XXXX-XXXX-XXXX")

    if st.button("Unlock"):
        ok, msg, row = verify_pin_and_bind(pin_input)
        if ok:
            st.session_state.auth_ok = True
            st.session_state.auth_row = row
            st.session_state.tier = tier_from_name(row["tier"])
            st.success(msg)
        else:
            st.session_state.auth_ok = False
            st.session_state.auth_row = None
            st.error(msg)

    if st.session_state.auth_ok:
        st.success(f"Tier: {st.session_state.tier.name}")
        st.write(f"Docs used (session): {st.session_state.docs_used}/{st.session_state.tier.max_docs}")
        st.write(f"Max pages: {st.session_state.tier.max_pages}")
        st.write(f"Max file size: {st.session_state.tier.max_mb}MB")
    else:
        st.warning("Locked. Enter a valid PIN.")
        st.stop()

    st.divider()
    st.subheader("👑 Admin (Owner only)")
    admin_pin = st.text_input("Admin PIN", type="password", help="This is only for generating/revoking pins.")
    is_admin = bool(admin_pin) and (admin_pin == OWNER_ADMIN_PIN)

    if is_admin:
        st.success("Admin unlocked ✅")

        gen_tier = st.selectbox("Generate tier", ["FREE", "PAID"], index=0)
        max_devices = st.selectbox("Max devices", [1, 2, 3], index=0)
        label = st.text_input("Label (e.g. 'friend varsity', 'order #123')")
        expiry_days = st.number_input("Expiry days (0 = no expiry)", min_value=0, max_value=365, value=30, step=1)

        if st.button("➕ Generate PIN"):
            raw_pin = generate_pin()
            payload = {
                "pin_hash": hash_pin(raw_pin),
                "tier": gen_tier,
                "label": label.strip() if label else None,
                "status": "ACTIVE",
                "max_devices": int(max_devices),
                "devices": [],
            }
            if int(expiry_days) > 0:
                exp = datetime.now(timezone.utc) + timedelta(days=int(expiry_days))
                payload["expires_at"] = exp.isoformat()

            sb().table("pins").insert(payload).execute()
            st.code(raw_pin)  # show ONCE

        st.markdown("### 🔎 Search pins")
        q = st.text_input("Search by label keyword", key="search_label")
        if st.button("Search"):
            data = sb().table("pins") \
                .select("id,tier,label,status,max_devices,devices,created_at,expires_at") \
                .ilike("label", f"%{q}%") \
                .order("created_at", desc=True) \
                .limit(50).execute().data
            st.session_state.admin_search = data

        data = st.session_state.get("admin_search", [])
        if data:
            for row in data:
                dev_count = len(row.get("devices") or [])
                st.write(f"**#{row['id']}** • {row['tier']} • {row.get('label','-')} • {row['status']} • devices {dev_count}/{row['max_devices']}")
                cols = st.columns(3)
                if cols[0].button("Revoke", key=f"rev_{row['id']}"):
                    sb().table("pins").update({"status": "REVOKED"}).eq("id", row["id"]).execute()
                    st.warning("Revoked.")
                if cols[1].button("Reset devices", key=f"rst_{row['id']}"):
                    sb().table("pins").update({"devices": []}).eq("id", row["id"]).execute()
                    st.info("Devices cleared.")
                if cols[2].button("Details", key=f"det_{row['id']}"):
                    st.json(row)
    else:
        st.caption("Admin tools locked.")


# =========================
# Main Settings
# =========================
tier = st.session_state.tier

left_settings, right_settings = st.columns([1, 1])

with left_settings:
    user_type = st.selectbox(
        "Who are you?",
        ["Varsity/College", "High school learner", "Primary learner", "Trader", "General"],
        index=0
    )

with right_settings:
    doc_mode = st.selectbox(
        "Document type",
        ["Notes/Book (PDF/DOCX/TXT)", "Slides (PDF only)"],
        index=0,
        help="Slides mode shows page-by-page highlights and is optimized for slide decks."
    )

summary_len = st.slider("Summary length (sentences)", 3, 12, 7)
keyword_count = st.slider("Keywords to highlight", 5, 30, 12)
question_count = st.slider("Questions to generate", 5, 25, 12)
flashcard_count = st.slider("Flashcards to generate", 5, 25, 10)

st.divider()

# =========================
# Upload (locked until PIN is valid)
# =========================
uploaded = st.file_uploader(
    f"Upload a document (Tier {tier.name} • {tier.max_docs} docs/session • Max {tier.max_mb}MB)",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=False
)

if not uploaded:
    st.stop()

if st.session_state.docs_used >= tier.max_docs:
    st.error("You reached your document limit for this session.")
    st.stop()

file_bytes = uploaded.getvalue()
if bytes_mb(file_bytes) > tier.max_mb:
    st.error(f"File too large: {bytes_mb(file_bytes):.1f}MB. Limit {tier.max_mb}MB.")
    st.stop()

filename = uploaded.name
ext = filename.lower().split(".")[-1]


# =========================
# PDF page range / section-only
# =========================
start_page, end_page = 1, 1
pdf_sections: List[Tuple[str, int, int]] = []
total_pages = None

if ext == "pdf":
    try:
        total_pages = pdf_total_pages(file_bytes)
    except Exception:
        total_pages = 0

    if total_pages <= 0:
        st.error("Could not read this PDF (might be corrupted or scanned-image only).")
        st.stop()

    st.subheader("⚡ PDF Processing Options")
    mode = st.radio(
        "Choose processing mode",
        ["Auto section (TOC → headings)", "Manual page range", "Full (limited by tier)"],
        horizontal=True
    )

    if mode == "Full (limited by tier)":
        start_page = 1
        end_page = min(total_pages, tier.max_pages)
        if total_pages > tier.max_pages:
            st.warning(f"Tier limit: processing first {end_page} pages (out of {total_pages}).")

    elif mode == "Manual page range":
        c1, c2 = st.columns(2)
        with c1:
            start_page = st.number_input("Start page", min_value=1, max_value=total_pages, value=1, step=1)
        with c2:
            end_page = st.number_input("End page", min_value=1, max_value=total_pages, value=min(10, total_pages), step=1)

        span = abs(int(end_page) - int(start_page)) + 1
        if span > tier.max_pages:
            st.error(f"Selected {span} pages, but your tier allows max {tier.max_pages}. Reduce the range.")
            st.stop()

    else:
        scan_pages = st.slider("If TOC is missing, scan headings in first N pages", 5, min(150, total_pages), 30)
        toc = toc_sections(file_bytes)
        if toc:
            headings = toc
            st.success("TOC found ✅")
        else:
            headings = scan_headings(file_bytes, scan_pages=scan_pages)
            if headings:
                st.info("TOC not found. Using scanned headings ✅")
            else:
                st.warning("No TOC/headings found. Use Manual page range or Full.")
                headings = [("Whole Document", 0)]

        pdf_ranges = heading_ranges(headings, total_pages)
        labels = [f"{t} (pages {s}-{e})" for (t, s, e) in pdf_ranges]
        pick = st.selectbox("Choose section", labels)
        title, start_page, end_page = pdf_ranges[labels.index(pick)]

        span = end_page - start_page + 1
        if span > tier.max_pages:
            st.warning(f"Section is {span} pages; trimming to tier max {tier.max_pages}.")
            end_page = start_page + tier.max_pages - 1

else:
    mode = "Non-PDF"


# =========================
# Process Button
# =========================
if st.button("⚡ Process Now", type="primary"):
    prog = st.progress(0, text="Starting...")

    def progress_cb(done: int, total: int):
        pct = int((done / max(1, total)) * 100)
        prog.progress(pct, text=f"Extracting… {done}/{total}")

    try:
        if ext == "pdf":
            text, per_page, pages_read = extract_pdf_range(file_bytes, int(start_page), int(end_page), progress_cb=progress_cb)
        elif ext == "docx":
            prog.progress(30, text="Reading DOCX…")
            text = extract_docx_text(file_bytes)
            per_page, pages_read = [], None
            prog.progress(100, text="Done.")
        else:
            prog.progress(30, text="Reading TXT…")
            text = extract_txt_text(file_bytes)
            per_page, pages_read = [], None
            prog.progress(100, text="Done.")

        if not text or len(text) < 120:
            prog.empty()
            st.error("No readable text extracted. If this is a scanned PDF (image-only), OCR is needed (we can add later).")
            st.stop()

        # analysis
        prog.progress(85, text="Analyzing…")
        keywords = extract_keywords(text, top_k=keyword_count)
        summary = summarize(text, n_sent=summary_len, user_type=user_type)
        takeaways = key_takeaways(text, k=min(10, max(5, summary_len)))
        questions = generate_questions(text, user_type=user_type, n=question_count)
        sections_text = split_into_sections(text)

        result = {
            "name": filename,
            "ext": ext,
            "mode": mode,
            "doc_mode": doc_mode,
            "user_type": user_type,
            "range": (int(start_page), int(end_page)) if ext == "pdf" else None,
            "pages_read": pages_read,
            "size_mb": bytes_mb(file_bytes),
            "keywords": keywords,
            "summary": summary,
            "takeaways": takeaways,
            "questions": questions,
            "text": text,
            "per_page": per_page if ext == "pdf" else [],
            "sections": sections_text,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }

        st.session_state.last_result = result
        st.session_state.history.insert(0, result)
        st.session_state.docs_used += 1

        prog.progress(100, text="Done ✅")
        time.sleep(0.2)
        prog.empty()

    except Exception as e:
        prog.empty()
        st.error(f"Processing failed: {type(e).__name__}: {e}")
        st.stop()


# =========================
# Show Results
# =========================
res = st.session_state.last_result
if not res:
    st.stop()

st.subheader("✅ Results")
meta = f"File: {res['name']} • Size: {res['size_mb']:.1f}MB • User type: {res['user_type']}"
if res.get("range"):
    meta += f" • Pages: {res['range'][0]}-{res['range'][1]}"
st.caption(meta)

tabs = st.tabs(["✅ Explanation", "🟣 Highlights", "❓ Questions", "📘 Glossary", "🃏 Flashcards", "🧩 Sections", "🆚 Compare", "🕘 History", "⬇️ Export"])

with tabs[0]:
    st.markdown("### Summary (readable)")
    st.write("\n".join([f"- {s}" for s in res["summary"]]))

    st.markdown("### Key Takeaways")
    for t in res["takeaways"]:
        st.markdown(f"- {t}")

with tabs[1]:
    st.markdown("### Keywords")
    st.write(", ".join(res["keywords"]) if res["keywords"] else "(none)")

    st.markdown("### Highlighted inside original text")
    if res["doc_mode"].startswith("Slides") and res["ext"] == "pdf":
        st.info("Slides mode: showing page-by-page highlights (faster + clearer for decks).")
        per = res.get("per_page") or []
        base_page = res["range"][0] if res.get("range") else 1
        for i, ptxt in enumerate(per, start=base_page):
            if not ptxt:
                continue
            st.markdown(f"#### Slide/Page {i}")
            kws_slide = extract_keywords(ptxt, top_k=min(10, keyword_count))
            st.markdown(highlight_html(ptxt, kws_slide, max_chars=25000), unsafe_allow_html=True)
    else:
        st.markdown(highlight_html(res["text"], res["keywords"]), unsafe_allow_html=True)

with tabs[2]:
    st.markdown("### Questions (non-repeating)")
    for i, q in enumerate(res["questions"], start=1):
        st.markdown(f"**{i}.** {q}")

with tabs[3]:
    st.markdown("### Glossary / Definitions")
    term = st.text_input("Type a word/term to search in the document", key="gloss_term")
    if st.button("Search term"):
        hits = glossary_search(term, res["text"], limit=8)
        if not hits:
            st.warning("No matches found in this document.")
        else:
            for h in hits:
                st.markdown(f"- {h}")

with tabs[4]:
    st.markdown("### Flashcards (practice mode)")
    if st.button("Generate flashcards"):
        st.session_state.cards = make_flashcards(res["text"], res["keywords"], n_cards=flashcard_count)
        st.session_state.card_answers = {}
        st.success(f"Generated {len(st.session_state.cards)} flashcards!")

    cards = st.session_state.cards
    if not cards:
        st.info("Generate flashcards to start.")
    else:
        for idx, c in enumerate(cards):
            st.markdown(f"#### Card {idx+1}")
            st.write(c["q"])
            if c["type"] == "mcq":
                choice = st.radio("Choose", c["options"], key=f"mcq_{idx}")
                st.session_state.card_answers[idx] = choice
            else:
                typed = st.text_input("Type your answer", key=f"typed_{idx}")
                st.session_state.card_answers[idx] = typed
            st.divider()

        if st.button("Check answers"):
            score = 0
            for idx, c in enumerate(cards):
                user_ans = (st.session_state.card_answers.get(idx) or "").strip()
                correct = (c["a"] or "").strip()
                ok = user_ans.lower() == correct.lower() if c["type"] == "mcq" else (correct.lower() in user_ans.lower())
                if ok:
                    score += 1
                    st.success(f"Card {idx+1}: Correct ✅")
                else:
                    st.error(f"Card {idx+1}: Wrong ❌ (Correct: {correct})")
            st.info(f"Score: {score}/{len(cards)}")

with tabs[5]:
    st.markdown("### Section Splitter (text-based)")
    sections = res.get("sections") or [("Whole Document", res["text"])]
    titles = [f"{i+1}. {t}" for i, (t, _) in enumerate(sections)]
    pick = st.selectbox("Choose a section", titles)
    idx = titles.index(pick)
    title, sec_text = sections[idx]

    st.markdown(f"#### {title}")
    sec_kws = extract_keywords(sec_text, top_k=min(12, keyword_count))
    sec_sum = summarize(sec_text, n_sent=min(summary_len, 8), user_type=res["user_type"])
    st.write("\n".join([f"- {s}" for s in sec_sum]) if sec_sum else "No summary.")
    st.markdown(highlight_html(sec_text, sec_kws, max_chars=60000), unsafe_allow_html=True)

with tabs[6]:
    st.markdown("### Compare (last two processed docs this session)")
    if len(st.session_state.history) < 2:
        st.info("Process at least 2 docs to compare.")
    else:
        a = st.session_state.history[0]
        b = st.session_state.history[1]
        c1, c2 = st.columns(2)
        with c1:
            st.markdown(f"#### A: {a['name']}")
            st.write("\n".join([f"- {s}" for s in a["summary"]]))
        with c2:
            st.markdown(f"#### B: {b['name']}")
            st.write("\n".join([f"- {s}" for s in b["summary"]]))

        common = sorted(set([k.lower() for k in a["keywords"]]) & set([k.lower() for k in b["keywords"]]))
        st.markdown("#### Common keywords")
        st.write(", ".join(common) if common else "(none)")

with tabs[7]:
    st.markdown("### History (this session)")
    if not st.session_state.history:
        st.info("No history yet.")
    else:
        for h in st.session_state.history[:10]:
            meta = f"{h['created_at']} • {h['user_type']} • {h['size_mb']:.1f}MB"
            if h.get("range"):
                meta += f" • pages {h['range'][0]}-{h['range'][1]}"
            st.markdown(f"**{h['name']}**")
            st.caption(meta)
            st.write(", ".join(h["keywords"][:12]) if h.get("keywords") else "")
            st.divider()

        export = {"exported_at": datetime.now().isoformat(timespec="seconds"), "history": st.session_state.history}
        st.download_button(
            "⬇️ Download History (JSON)",
            data=json.dumps(export, indent=2),
            file_name="doculite_history.json",
            mime="application/json"
        )

with tabs[8]:
    st.markdown("### Export Summary + Questions")
    pdf_bytes = export_pdf(APP_NAME, res["summary"], res["takeaways"], res["questions"])
    docx_bytes = export_docx(APP_NAME, res["summary"], res["takeaways"], res["questions"])

    st.download_button("⬇️ Download PDF", data=pdf_bytes, file_name="doculite_export.pdf", mime="application/pdf")
    st.download_button("⬇️ Download Word (DOCX)", data=docx_bytes, file_name="doculite_export.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.caption("Note: scanned/image PDFs need OCR to extract text. We can add OCR later (it costs CPU/time).")
